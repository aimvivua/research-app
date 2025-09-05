import io
import pandas as pd
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

def generate_docx_report(title, notes, df=None):
    """
    Generates a DOCX report with a title, notes, and an optional data table.
    """
    document = Document()
    
    document.add_heading(title, 0)
    document.add_paragraph(notes)
    
    if df is not None and not df.empty:
        document.add_heading('Dataset Preview', level=1)
        
        # Convert DataFrame to a list of lists for the Word table
        table_data = [df.columns.tolist()] + df.values.tolist()
        table = document.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(df.columns):
            hdr_cells[i].text = str(header)
            
        # Add data
        for row_data in df.itertuples(index=False):
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf_report(title, notes, df=None):
    """
    Generates a PDF report with a title, notes, and an optional data table.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    
    styles = getSampleStyleSheet()
    elements = []

    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=18, spaceAfter=20)
    elements.append(Paragraph(title, title_style))
    
    # Notes
    elements.append(Paragraph(notes, styles['Normal']))
    elements.append(Spacer(1, 0.25 * inch))
    
    if df is not None and not df.empty:
        elements.append(Paragraph("Dataset Preview", styles['Heading2']))
        elements.append(Spacer(1, 0.1 * inch))
        
        # Convert DataFrame to a list of lists for the PDF table
        table_data = [df.columns.tolist()] + df.values.tolist()
        
        # Create the table style
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ])
        
        # Create and add the table
        table = Table(table_data)
        table.setStyle(table_style)
        elements.append(table)

    doc.build(elements)
    buffer.seek(0)
    return buffer